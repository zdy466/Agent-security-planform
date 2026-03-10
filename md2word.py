"""Convert Markdown to Word document"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

doc.add_heading('AgentShield OS：企业级 AI Agent 安全守护者', 0)

doc.add_paragraph(
    '在 AI Agent 广泛应用的时代，您是否考虑过它的安全问题？\n\n'
    '随着 ChatGPT、Claude 等大语言模型的兴起，AI Agent 正在成为企业数字化转型的重要工具。'
    '它们可以代替员工完成数据分析、客户服务、内容生成等复杂任务。\n\n'
    '然而，您是否知道：\n'
    '• 🔓 敏感数据泄露：AI Agent 可能在输出中不经意暴露客户个人信息\n'
    '• 🎯 Prompt 注入：黑客可通过恶意指令操控 AI Agent 执行非授权操作\n'
    '• 🔧 工具滥用：AI Agent 调用外部工具时可能造成数据外泄\n'
    '• ⚠️ 合规风险：AI 输出可能违反 GDPR、HIPAA 等法规\n\n'
    '您的企业准备好应对这些挑战了吗？'
)

doc.add_heading('🛡️ AgentShield OS — 您的 AI 安全盾牌', level=1)

doc.add_paragraph(
    'AgentShield OS 是中国企业自主研发的企业级 AI Agent 安全运行层，为 AI Agent 提供全面的安全保护。'
)

doc.add_heading('核心能力', level=2)

table = doc.add_table(rows=9, cols=2)
table.style = 'Light Grid Accent 1'

headers = table.rows[0].cells
headers[0].text = '安全模块'
headers[1].text = '功能说明'

data = [
    ('LLM Data Firewall', '智能识别并阻止敏感数据泄露，支持邮箱、电话、身份证、银行卡等'),
    ('Prompt Injection Firewall', '实时检测并拦截恶意指令注入攻击'),
    ('Tool Guard', '工具白名单、参数验证、沙盒执行，全面管控 AI 工具调用'),
    ('Data Gateway', '字段级/行级权限控制，确保数据访问安全'),
    ('Behavior Monitor', 'AI 行为分析，及时发现异常操作'),
    ('Policy Engine', '统一策略管理，一键配置安全规则'),
    ('Compliance Manager', '内置 GDPR、HIPAA、SOC2 合规模块'),
    ('Attack Simulation', '定期模拟攻击测试，验证安全防护能力'),
]

for i, (module, desc) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = module
    row[1].text = desc

doc.add_heading('适用场景', level=1)

scenarios = [
    ('金融行业', '防止 AI 泄露客户账户信息、满足金融合规要求、保护交易数据安全'),
    ('医疗健康', '防止患者隐私数据外泄、符合 HIPAA 合规、保障医疗 AI 安全运行'),
    ('电商零售', '保护用户信息安全、防止恶意指令干扰业务、提升客户信任度'),
    ('企业通用', '内部 AI 助手安全部署、多 Agent 协作环境管控、AI 应用安全审计'),
]

for title, content in scenarios:
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)

doc.add_heading('为什么选择 AgentShield OS？', level=1)

reasons = [
    ('自主可控', '100% 自主研发、源代码开放透明、持续安全更新'),
    ('易于集成', '支持 LangChain、AutoGPT 等主流框架、Python SDK 开箱即用、灵活配置即插即用'),
    ('全面防护', '覆盖数据安全/工具安全/行为安全、实时监控与告警、完整审计日志'),
    ('合规内置', '内置 GDPR/HIPAA/SOC2 合规规则、自动生成合规报告、降低企业合规成本'),
]

for title, content in reasons:
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)

doc.add_heading('快速开始', level=1)

code = '''# 安装
pip install git+https://github.com/zdy466/Agent-security-planform.git

# 使用
from agentshield import AgentShieldClient

client = AgentShieldClient()

# AI 输入安全检查
result = client.process_input("请帮我查询用户邮箱: john@example.com")
# {'allowed': False, 'reason': 'sensitive_data_detected'}'''

p = doc.add_paragraph(code)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)

doc.add_heading('联系我们', level=1)

doc.add_paragraph(
    '• 🌐 GitHub: https://github.com/zdy466/Agent-security-planform\n'
    '• 📧 邮箱: 18335057001@163.com'
)

doc.add_heading('开源协议', level=1)

doc.add_paragraph('MIT License — 免费商用，欢迎贡献！')

doc.add_paragraph('')
final = doc.add_paragraph('让 AI Agent 安全地为您服务')
final.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'd:\security\AgentShield_OS_宣传文章.docx')
print("Word 文档已保存到: d:\\security\\AgentShield_OS_宣传文章.docx")
